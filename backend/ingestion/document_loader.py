# ingestion/document_loader.py

import fitz  # PyMuPDF
import docx
import os
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass

@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]

class DocumentLoader:
    """Loads PDF, DOCX, TXT, and .py files into Document objects."""

    SUPPORTED = {".pdf", ".docx", ".txt", ".py", ".md"}

    def load_directory(self, dir_path: str) -> List[Document]:
        docs = []
        for path in Path(dir_path).rglob("*"):
            if path.suffix in self.SUPPORTED:
                docs.extend(self.load_file(str(path)))
        return docs

    def load_file(self, file_path: str) -> List[Document]:
        ext = Path(file_path).suffix.lower()
        loaders = {
            ".pdf":  self._load_pdf,
            ".docx": self._load_docx,
            ".txt":  self._load_text,
            ".py":   self._load_text,
            ".md":   self._load_text,
        }
        loader = loaders.get(ext)
        if not loader:
            raise ValueError(f"Unsupported file type: {ext}")
        return loader(file_path)

    def _load_pdf(self, path: str) -> List[Document]:
        docs = []
        pdf = fitz.open(path)
        for i, page in enumerate(pdf):
            text = page.get_text().strip()
            if text:
                docs.append(Document(
                    content=text,
                    metadata={"source": path, "page": i + 1, "type": "pdf"}
                ))
        return docs

    def _load_docx(self, path: str) -> List[Document]:
        doc = docx.Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(content=text, metadata={"source": path, "type": "docx"})]

    def _load_text(self, path: str) -> List[Document]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        file_type = Path(path).suffix.lstrip(".")
        return [Document(content=text, metadata={"source": path, "type": file_type})]